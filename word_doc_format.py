from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

def gerar_documento_word_com_formatacao():
    # Criando um novo documento
    doc = Document()

    # Adicionando um título
    titulo = doc.add_heading(level=1)
    titulo_run = titulo.add_run("Exemplo de Documento Word com Formatação")
    titulo_run.bold = True

    # Adicionando um parágrafo
    paragrafo = doc.add_paragraph("Este é um parágrafo de exemplo.")
    paragrafo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Adicionando um parágrafo com formatação diferente
    paragrafo2 = doc.add_paragraph("Este parágrafo está alinhado à direita.")
    paragrafo2.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    # Adicionando uma lista
    lista = doc.add_paragraph()
    lista.add_run("Lista de itens:").bold = True
    itens = ["Item 1", "Item 2", "Item 3"]
    for item in itens:
        lista.add_run("\n\u2022 " + item)

    # Salvando o documento
    doc.save("documento_word_com_formatacao.docx")

if __name__ == "__main__":
    gerar_documento_word_com_formatacao()